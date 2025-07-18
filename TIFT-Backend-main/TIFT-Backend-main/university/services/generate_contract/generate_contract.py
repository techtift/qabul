import os
import tempfile
import subprocess
from datetime import datetime
import re
import fitz 
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.table import Table
import qrcode
from university.models.application import StudentApplication
from core.exceptions.exception import CustomApiException
from core.exceptions.error_messages import ErrorCodes
from user.models.student import Student
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL
from django.utils.translation import gettext as _
from university.services.study_duration.get_study_duration import get_study_duration

TEMPLATE_PATH_UZ = "contract_2025_uz.docx"
TEMPLATE_PATH_RU = "contract_2025_ru.docx"
STAMP_IMAGE_PATH = "pechat.png" 
HEADER_IMAGE_PATH = "certificate_header.jpg"  # Path to the header image

UZ_MONTHS = {
    1: "Yanvar", 2: "Fevral", 3: "Mart", 4: "Aprel",
    5: "May", 6: "Iyun", 7: "Iyul", 8: "Avgust",
    9: "Sentyabr", 10: "Oktyabr", 11: "Noyabr", 12: "Dekabr"
}


# === 🔁 FORMATNI SAQLAB, PLACEHOLDERLARNI oddiy matn bilan almashtirish ===
def replace_placeholders(doc: Document, context: dict):
    def replace_in_paragraph(paragraph: Paragraph):
        full_text = paragraph.text
        for key, value in context.items():
            placeholder = f"{{{key}}}"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
        
        # Paragraphni tozalab, yangi `run` qo‘shish
        if paragraph.text != full_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text

    def replace_in_cell(cell):
        for paragraph in cell.paragraphs:
            replace_in_paragraph(paragraph)
        for table in cell.tables:
            replace_in_table(table)

    def replace_in_table(table: Table):
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        replace_in_table(table)
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for paragraph in header_footer.paragraphs:
                replace_in_paragraph(paragraph)
            for table in header_footer.tables:
                replace_in_table(table)


# === 📎 PECHAT + QR NI OXIRGI SAHIFAGA JOYLASHTIRISH ===
def overlay_stamp_on_pdf(pdf_path: str, stamp_path: str, output_path: str, qr_path: str, language="uz"):
    doc = fitz.open(pdf_path)
    last_page = doc[-1]

    # 📌 Koordinatalar (test orqali moslashtiring)
    if language == "uz":
        stamp_rect = fitz.Rect(80, 400, 300, 620)
        qr_rect = fitz.Rect(80, 620, 180, 740)  # jadvalning chap pastki qismi
    else:
        stamp_rect = fitz.Rect(80, 450, 300, 670)
        qr_rect = fitz.Rect(80, 650, 180, 770)

    last_page.insert_image(stamp_rect, filename=stamp_path, overlay=True)
    last_page.insert_image(qr_rect, filename=qr_path, overlay=True)

    doc.save(output_path)
    doc.close()


# === QR KOD GENERATSIYA QILISH ===
def generate_qr_code(url: str) -> str:
    tmp_qr_path = tempfile.NamedTemporaryFile(suffix=".png", delete=False).name
    qr = qrcode.make(url)
    qr.save(tmp_qr_path)
    return tmp_qr_path


# === FAYL NOMINI TOZALASH ===
def safe_filename(name):
    return re.sub(r'[^a-zA-Z0-9_]', '_', name)


# === PDF KONTRAKT YARATISH ===
def generate_contract_pdf(context, qr_url, language="ru"):
    doc = Document(TEMPLATE_PATH_UZ if language == "uz" else TEMPLATE_PATH_RU)
    replace_placeholders(doc, context)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        temp_docx_path = tmp_docx.name
        doc.save(temp_docx_path)

    # .docx → .pdf
    pdf_dir = tempfile.gettempdir()
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', pdf_dir, temp_docx_path
    ], check=True)

    base_filename = os.path.splitext(os.path.basename(temp_docx_path))[0]
    pdf_path = os.path.join(pdf_dir, base_filename + ".pdf")
    stamped_pdf_path = os.path.join(pdf_dir, f"{base_filename}_stamped.pdf")

    # ✅ QR rasmni yaratamiz
    qr_path = generate_qr_code(qr_url)

    # ✅ Pechat va QR ni PDF oxiriga qo'yamiz
    overlay_stamp_on_pdf(pdf_path, STAMP_IMAGE_PATH, stamped_pdf_path, qr_path, language=language)

    with open(stamped_pdf_path, 'rb') as f:
        pdf_data = f.read()

    # Vaqtinchalik fayllarni o‘chiramiz
    os.remove(temp_docx_path)
    os.remove(pdf_path)
    os.remove(stamped_pdf_path)
    os.remove(qr_path)

    full_name = context.get("full_name", "contract").replace(" ", "_")
    safe_name = safe_filename(full_name)
    pdf_filename = f"{safe_name}.pdf"

    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
    return response


# === STUDENT UCHUN KONTRAKT YASASH ===
def generate_student_contract(student_id, language="uz"):
    student = Student.objects.filter(id=student_id).first()
    if not student:
        raise CustomApiException(ErrorCodes.NOT_FOUND, "Student not found")

    student_application = StudentApplication.objects.filter(student=student).first()
    if not student_application:
        raise CustomApiException(ErrorCodes.NOT_FOUND, "Student application not found")

    def get_faculty_price(faculty, s_type):
        if s_type and s_type.id == 1:
            return faculty.day_price
        elif s_type and s_type.id == 2:
            return faculty.night_price
        else:
            return faculty.external_price

    program = student_application.program
    faculty = student_application.faculty
    study_type = student_application.study_type
    created_date = datetime.now()

    full_name = f"{student.first_name} {student.last_name} {student.father_name}"
    passport_number = student.passport_number or "N/A"
    study_language = dict(StudentApplication._meta.get_field('lang').choices).get(student_application.lang, "Uzbek")
    study_duration = get_study_duration(faculty.id, study_type.id) if faculty and study_type else None
    course = f"{student_application.transfer_level}-kurs" if student_application.is_transfer else "1-kurs"
    price = get_faculty_price(faculty, study_type)

    if study_duration:
        result = f"{study_duration} yil"
    elif program and program.id == 29:
        result = "2 yil"
    elif study_type and study_type.id == 2:
        result = "4.5 yil"
    elif study_type and study_type.id == 1:
        result = "4 yil"
    else:
        result = "5 yil"

    context = {
        "contract_id": datetime.now().strftime("%Y") + "_" + str(student.passport_number),
        "contract_date": created_date.strftime("%d"),
        "contract_month": UZ_MONTHS[created_date.month],
        "full_name": full_name,
        "passport_number": passport_number,
        "birth_date": student.birth_date.strftime("%d.%m.%Y") if student.birth_date else "N/A",
        "phone_number": student.user.phone_number if student.user else "N/A",
        "program": program.name if program else "N/A",
        "study_type": study_type.name if study_type else "N/A",
        "study_duration": result,
        "address": student.address or "N/A",
        "course": course,
        "study_language": study_language,
        "faculty": faculty.name if faculty else "N/A",
        "price": price,
    }

    qr_url = f"https://qabul.tift.uz/api/v1/university/generate-contract/?student_id={student.id}"
    return generate_contract_pdf(context, qr_url, language=language)

def insert_blank_paragraphs_after(paragraph, count=2):
    for _ in range(count):
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)


def generate_student_certificate(student_id, language="uz"):
    student = Student.objects.filter(id=student_id).first()
    if not student:
        raise CustomApiException(ErrorCodes.NOT_FOUND, "Student not found")

    student_application = StudentApplication.objects.filter(student=student).first()
    if not student_application:
        raise CustomApiException(ErrorCodes.NOT_FOUND, "Student application not found")

    program = student_application.program
    faculty = student_application.faculty
    course = f"{student_application.transfer_level}" if student_application.is_transfer else "1"
    full_name = f"{student.first_name} {student.last_name} {student.father_name}"

    context = {
        "id": student.id,
        "given_date": datetime.now().strftime("%d.%m.%Y"),
        "name": full_name,
        "program_name": program.name if program else "N/A",
        "faculty_name": faculty.name if faculty else "N/A",
        "course": course,
    }

    template_path = "certificate_2025_uz.docx" if language == "uz" else "certificate_2025_ru.docx"
    doc = Document(template_path)

    # Reduce top margin
    for section in doc.sections:
        section.top_margin = Inches(0.2)

    # Replace placeholders
    replace_placeholders(doc, context)

    # Insert header image
    inserted = False
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(HEADER_IMAGE_PATH, width=Inches(8))
            doc._body._element.remove(p._p)
            doc._body._element.insert(idx, p._p)
            inserted = True
            break
    if not inserted:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(HEADER_IMAGE_PATH, width=Inches(6))
        doc._body._element.insert(0, p._p)

    # Add 2 blank paragraphs after "certificate number"
    for paragraph in doc.paragraphs:
        if "certificate number" in paragraph.text.lower():
            insert_blank_paragraphs_after(paragraph, 1)
            break

    # Remove empty paragraphs before the table
    for paragraph in reversed(doc.paragraphs):
        if paragraph.text.strip() == "":
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
        else:
            break  # stop when we hit a non-empty paragraph

    # Remove existing 'Rektor ... A.Nodirov' paragraph if exists
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if (language == "uz" and text.startswith("Rektor") and "A.Nodirov" in text) or \
           (language == "ru" and text.startswith("Ректор") and "А.Нодиров" in text):
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
            break

    # Create stamp + text table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    row.height = Inches(1.0)

    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Left: Rektor
    left_para = row.cells[0].paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_para.add_run("Rektor" if language == "uz" else "Ректор")

    # Center: Stamp image
    center_para = row.cells[1].paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_para.add_run().add_picture(STAMP_IMAGE_PATH, width=Inches(2.0))

    # Right: A.Nodirov
    right_para = row.cells[2].paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right_para.add_run("A.Nodirov" if language == "uz" else "А.Нодиров")

    # Save temporary DOCX
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        temp_docx_path = tmp_docx.name
        doc.save(temp_docx_path)

    # Convert DOCX -> PDF
    pdf_dir = tempfile.gettempdir()
    subprocess.run([
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', pdf_dir, temp_docx_path
    ], check=True)

    pdf_path = os.path.join(pdf_dir, os.path.splitext(os.path.basename(temp_docx_path))[0] + ".pdf")
    with open(pdf_path, 'rb') as f:
        pdf_data = f.read()

    os.remove(temp_docx_path)
    os.remove(pdf_path)

    safe_name = safe_filename(full_name)
    pdf_filename = f"malumotnoma_{safe_name}.pdf"

    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
    return response